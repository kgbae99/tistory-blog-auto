"""프로젝트 보고서 Word 파일 생성."""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r+1].cells[c].text = str(val)
    doc.add_paragraph("")


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Malgun Gothic"
style.font.size = Pt(10)

# 표지
for _ in range(3):
    doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("건강온도사(행복++)")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0xE8, 0x7A, 0x8A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("티스토리 블로그 자동화 시스템\n프로젝트 구축 보고서")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026년 3월 15일\ngithub.com/kgbae99/tistory-blog-auto\nkgbae2369.tistory.com")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
doc.add_page_break()

# 1. 프로젝트 개요
doc.add_heading("1. 프로젝트 개요", level=1)
doc.add_paragraph("티스토리 블로그 수익화를 위한 완전 자동화 시스템.")
add_table(doc, ["항목", "내용"], [
    ["블로그", "건강온도사(행복++) - kgbae2369.tistory.com"],
    ["주제", "건강/생활/뷰티/복지 정보"],
    ["대상", "20~60대"],
    ["수익모델", "쿠팡 파트너스 + 구글 애드센스"],
    ["콘텐츠 엔진", "GPT-5 Mini (Gemini 폴백)"],
    ["자동화", "GitHub Actions 매일 07:00 KST"],
    ["총 포스트", "484개 + 매일 3개 자동 생성"],
])

# 2. 시스템 아키텍처
doc.add_heading("2. 시스템 아키텍처", level=1)
doc.add_heading("핵심 모듈", level=2)
add_table(doc, ["모듈", "기능"], [
    ["src/content/generator.py", "GPT-5 Mini / Gemini 듀얼 콘텐츠 생성"],
    ["src/content/dedup_checker.py", "제목/키워드 중복 방지 (Jaccard 50%)"],
    ["src/content/image_search.py", "GitHub 고정 이미지 26개 카테고리별 매칭"],
    ["src/content/internal_links.py", "블로그 크롤링 DB 기반 내부 링크 자동 매칭"],
    ["src/content/realtime_trends.py", "구글/네이버 트렌드 + 카테고리 균형 배분"],
    ["src/content/keyword_research.py", "자동완성 기반 롱테일 키워드 발굴"],
    ["src/coupang/api_client.py", "쿠팡 파트너스 HMAC API 연동"],
    ["src/coupang/smart_matcher.py", "30+ 카테고리 고수익 상품 자동 매핑"],
    ["src/adsense/ad_inserter.py", "애드센스 광고 3개 자동 삽입 (슬롯 ID 지원)"],
    ["src/seo/search_console.py", "Google Search Console API (색인 진단+요청)"],
    ["src/analytics/dashboard.py", "수익 대시보드 + 리포트"],
    ["src/analytics/traffic_analyzer.py", "유입 분석 → 콘텐츠 갭 → 트렌드 최신화"],
    ["src/notify/telegram.py", "텔레그램 알림 (선택)"],
])

doc.add_heading("스크립트", level=2)
add_table(doc, ["스크립트", "기능"], [
    ["scripts/generate_daily_posts.py", "매일 3개 포스트 자동 생성 (메인)"],
    ["scripts/crawl_blog_posts.py", "블로그 전체 크롤링 → 내부링크 DB"],
    ["scripts/request_indexing.py", "미색인 페이지 자동 색인 요청"],
])
doc.add_page_break()

# 3. 콘텐츠 생성
doc.add_heading("3. 콘텐츠 생성 엔진", level=1)
doc.add_paragraph("GPT-5 Mini 우선 사용, 실패 시 Gemini 2.5 Flash 자동 폴백.")
add_table(doc, ["항목", "GPT-5 Mini", "Gemini 2.5 Flash"], [
    ["역할", "기본 엔진", "폴백 엔진"],
    ["입력 가격", "$0.25/1M 토큰", "$0.30/1M 토큰"],
    ["출력 가격", "$2.00/1M 토큰", "$2.50/1M 토큰"],
    ["월 예상", "약 1,100원", "약 1,400원"],
])

doc.add_heading("포스트 HTML 구조", level=2)
for item in [
    "쿠팡 고지문 + 건강 면책문 (상단)",
    "대표 이미지 (GitHub 고정 URL)",
    "H1 제목 + 도입부 + 목차 (바로가기 없음)",
    "H2 섹션 1 → 애드센스 #1 (상단 반응형)",
    "H2 섹션 2~3 → 애드센스 #2 (중간 인피드)",
    "핵심 요약 카드 5개 + FAQ 3개",
    "애드센스 #3 (하단 반응형)",
    "마무리 + 쿠팡 추천 상품 3개 + 태그",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_page_break()

# 4. 수익화
doc.add_heading("4. 수익화 시스템", level=1)
doc.add_heading("쿠팡 파트너스", level=2)
add_table(doc, ["항목", "내용"], [
    ["인증", "HMAC-SHA256"],
    ["상품 수", "포스트당 3개 (이미지+가격+버튼)"],
    ["매칭", "키워드 → 30+ 카테고리 고수익 상품"],
])
doc.add_heading("구글 애드센스", level=2)
add_table(doc, ["항목", "내용"], [
    ["퍼블리셔", "ca-pub-5391209329273834"],
    ["상단 슬롯", "8208542949 (반응형)"],
    ["중간 슬롯", "8134684468 (인피드)"],
    ["하단 슬롯", "9545675342 (반응형)"],
    ["티스토리 설정", "자동광고 ON + 본문 상단/하단 ON"],
])

# 5. SEO
doc.add_heading("5. SEO 최적화", level=1)
add_table(doc, ["항목", "구현"], [
    ["메타 태그", "description, keywords 자동 삽입"],
    ["OG 태그", "og:title, og:description, og:type"],
    ["JSON-LD", "Article + FAQPage 스키마"],
    ["이미지", "반응형 + alt 빈값 + border-radius"],
    ["내부 링크", "크롤링 DB 49개 기반 자동 2개 삽입"],
])

# 6. Search Console
doc.add_heading("6. Google Search Console 연동", level=1)
add_table(doc, ["항목", "내용"], [
    ["인증", "OAuth2 (토큰 자동 갱신)"],
    ["검색 실적", "클릭/노출/CTR/순위 조회"],
    ["색인 검사", "URL Inspection API"],
    ["색인 요청", "Indexing API (미색인 자동 요청)"],
    ["초기 결과", "50개 중 43개 색인, 7개 요청 완료"],
    ["네이버", "서치어드바이저 등록 + 사이트맵/RSS 제출"],
])
doc.add_page_break()

# 7. 이미지
doc.add_heading("7. 이미지 관리", level=1)
add_table(doc, ["항목", "내용"], [
    ["저장", "GitHub assets/images/ 영구 저장 (26개)"],
    ["URL", "GitHub raw URL (절대 변동 없음)"],
    ["카테고리", "건강/음식/운동/봄/피곤/수면/알레르기/미세먼지/영양제/피부/다이어트"],
    ["중복 방지", "사용 이력 추적 + 최소 사용 우선 + pop 방식 원천 차단"],
    ["alt 텍스트", "빈값 처리 (불필요 텍스트 노출 방지)"],
])

# 8. 중복 방지
doc.add_heading("8. 중복 방지 시스템", level=1)
add_table(doc, ["항목", "방식"], [
    ["제목 중복", "Jaccard 유사도 50% 이상 → 중복 감지 + 스킵"],
    ["키워드 중복", "최근 30개 내 사용 키워드 자동 스킵"],
    ["이미지 중복", "사용 이력 DB + 2회 이상 사용 방지"],
    ["카테고리 분산", "건강/음식/뷰티 3개 매일 로테이션 (편중 방지)"],
    ["키워드 풀", "카테고리당 15개 = 총 45개 (15일 주기)"],
])

# 9~11 간략
doc.add_heading("9. 내부 링크 자동화", level=1)
doc.add_paragraph("블로그 484개 포스트 사이트맵 크롤링 → 49개 상세 인덱스 구축 → 키워드 매칭 점수 기반 관련 포스트 2개 자동 삽입")

doc.add_heading("10. 키워드 리서치", level=1)
doc.add_paragraph("네이버/구글 자동완성 기반 롱테일 확장. Search Console 유입 데이터 → 콘텐츠 갭 발견 → 추천 키워드 자동 반영.")

doc.add_heading("11. 수익 대시보드", level=1)
doc.add_paragraph("포스트 DB 등록/조회, 일일 수익 기록, 카테고리별 분석, 유입 기반 트렌드 최신화.")
doc.add_page_break()

# 12. GitHub Actions
doc.add_heading("12. GitHub Actions 자동화", level=1)
add_table(doc, ["항목", "내용"], [
    ["실행 시간", "매일 07:00 KST (UTC 22:00)"],
    ["포스트", "3개/일 (카테고리별 1개)"],
    ["색인 요청", "최근 20개 자동 검사"],
    ["저장소", "github.com/kgbae99/tistory-blog-auto"],
])

# 13. 환경 설정
doc.add_heading("13. 환경 설정 (GitHub Secrets 11개)", level=1)
add_table(doc, ["Secret", "용도"], [
    ["OPENAI_API_KEY", "GPT-5 Mini"],
    ["GEMINI_API_KEY", "Gemini 폴백"],
    ["COUPANG_ACCESS_KEY", "쿠팡 API"],
    ["COUPANG_SECRET_KEY", "쿠팡 HMAC"],
    ["ADSENSE_PUB_ID", "애드센스 퍼블리셔"],
    ["ADSENSE_SLOT_TOP/MID/BOTTOM", "광고 슬롯 3개"],
    ["KAKAO_EMAIL/PASSWORD", "티스토리 로그인"],
    ["GSC_TOKEN_JSON", "Search Console 토큰"],
])

# 14. 명령어
doc.add_heading("14. 명령어 가이드", level=1)
add_table(doc, ["기능", "명령어"], [
    ["포스트 3개 생성", "python scripts/generate_daily_posts.py"],
    ["블로그 크롤링", "python scripts/crawl_blog_posts.py --count 100"],
    ["색인 요청", "python scripts/request_indexing.py --count 30"],
    ["Search Console 리포트", 'python -c "from src.seo.search_console import print_full_report; print_full_report()"'],
    ["수익 대시보드", 'python -c "from src.analytics.dashboard import print_dashboard; print_dashboard()"'],
    ["유입 분석", 'python -c "from src.analytics.traffic_analyzer import print_traffic_report; print_traffic_report()"'],
])

# 15. 로드맵
doc.add_heading("15. 향후 로드맵", level=1)
add_table(doc, ["기간", "목표"], [
    ["1주차", "매일 3개 발행 → 주 21개 축적"],
    ["2~4주", "색인 200+ → 검색 노출 시작"],
    ["1~2개월", "일 방문자 100+ → 첫 수익"],
    ["3개월", "일 방문자 500+ → 월 5~10만원"],
])

doc.add_paragraph("")
doc.add_heading("매일 운영 루틴", level=2)
for r in [
    "07:00 - GitHub Actions 자동 실행 (포스트 3개 + 색인 요청)",
    "07:05 - GitHub output/posts/ 폴더 확인",
    "07:10 - HTML 파일 열기 → 복사 → 티스토리 HTML 모드 붙여넣기 (3회)",
    "07:20 - 발행 완료",
]:
    doc.add_paragraph(r, style="List Bullet")

# 저장
filepath = "output/건강온도사_블로그자동화_프로젝트보고서.docx"
doc.save(filepath)
print(f"Word 파일 생성 완료: {filepath}")
